import os
import numpy as np
from math import ceil

from docx import Document
from docx.enum.text import WD_COLOR_INDEX

WCI = {-1: WD_COLOR_INDEX.RED,
		0: WD_COLOR_INDEX.YELLOW, \
		1: WD_COLOR_INDEX.BRIGHT_GREEN, \
		2: WD_COLOR_INDEX.GRAY_25, \
		3: WD_COLOR_INDEX.TEAL}

def set_global_voc_cost(c):
	global GOLBAL_VOC_COST
	GOLBAL_VOC_COST = c

def log_star(x):
	"""
	Universal code length

	"""
	return 2 * ceil(np.log2(x)) + 1 if x != 0 else 0

def word_cost():
	return GOLBAL_VOC_COST

def sequence_cost(seq):
	"""
	Output encoding cost for a given sequence

	"""
	return len(seq) * word_cost()

def output_word(temp, cond, word_path):
	"""
	Output highlight content with office word document

	"""

	### Initialize document
	doc = Document()
	proc = doc.add_paragraph()
	for s, c in zip(['Slot', 'Matched', 'Substitution', 'Deletion', 'Insertion'], WCI.values()):
		font = proc.add_run(s).font
		font.highlight_color = c
		proc.add_run(' ')

	### Template content
	proc = doc.add_paragraph()
	proc.add_run('Template: \n')
	proc.add_run(temp.seq())
	proc.add_run('\n\n-----------------------------------------------------------------\n')

	### Iterate all aligned sequences
	for cs in cond:
		proc = doc.add_paragraph()
		for c, s in cs:
			font = proc.add_run(s).font
			font.highlight_color = WCI[c]
			proc.add_run(' ')

	doc.save(word_path)

def output_results(temp_arr, cond_arr, file_name, folder_name='results', html_name='graph.html', word_name='text.docx'):
	"""
	Output template results

	"""

	directory = os.path.join(folder_name, file_name)
	if not os.path.exists(directory):
		os.makedirs(directory)

	### Iterate all templates
	for idx, (temp, cond) in enumerate(zip(temp_arr, cond_arr)):
		temp_path = os.path.join(directory, 'template_' + str(idx + 1))
		if not os.path.exists(temp_path):
			os.makedirs(temp_path)

		### Output html
		temp.htmlOutput(open(os.path.join(temp_path, html_name), 'w'))

		### Output word document
		output_word(temp, cond, os.path.join(temp_path, word_name))
